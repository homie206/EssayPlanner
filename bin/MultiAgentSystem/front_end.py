import importlib.util
import os
import re
import time
from io import BytesIO
from pathlib import Path

import requests
import streamlit as st
from docx import Document
from fpdf import FPDF
from fpdf.errors import FPDFException

# Load agent_config.py directly (by file path) to avoid triggering
# bin/RichUI/__init__.py, which imports heavy orchestrator/agent deps.
_spec = importlib.util.spec_from_file_location(
    "agent_config",
    Path(__file__).resolve().parents[1] / "RichUI" / "agent_config.py",
)
_agent_config = importlib.util.module_from_spec(_spec)
_spec.loader.exec_module(_agent_config)
AGENT_CONFIG = _agent_config.AGENT_CONFIG


st.set_page_config(page_title="MultiAgentSystem UI", layout="centered")


# ----------------------------
# Helpers
# ----------------------------

YES_NO_MARKER = "[YES_NO]"

# Build reverse lookup: state_key -> agent_name
_STATE_KEY_TO_AGENT = {}
for _name, _cfg in AGENT_CONFIG.items():
    _STATE_KEY_TO_AGENT[_cfg["state_key"]] = _name
    # Also map the _reply key for idea_structurer (its state_key is idea_board)
    _STATE_KEY_TO_AGENT[f"{_name}_reply"] = _name


def post_json(url: str, payload: dict, timeout: int = 300) -> dict:
    r = requests.post(url, json=payload, timeout=timeout)
    r.raise_for_status()
    return r.json()


def role_from_key(node: str, key: str) -> str:
    """Map state keys / nodes to agent names using AGENT_CONFIG."""
    if key in _STATE_KEY_TO_AGENT:
        return _STATE_KEY_TO_AGENT[key]
    return node or "system"


def should_display(key: str, value) -> bool:
    """Only show agent replies (NOT the idea_board raw value)."""
    if value is None:
        return False
    if isinstance(value, str):
        return value.strip() != "" and key.endswith("_reply")
    return False


def render_chat_message(role: str, content: str):
    """Render a single chat message using st.chat_message()."""
    if role == "user":
        with st.chat_message("user", avatar="🧑‍🎓"):
            st.markdown(content)
    else:
        cfg = AGENT_CONFIG.get(role)
        if cfg:
            with st.chat_message(name=cfg["label"], avatar=cfg["emoji"]):
                st.markdown(
                    f'<div style="'
                    f'background:{cfg["bg_color"]};'
                    f'border-left:4px solid {cfg["st_color"]};'
                    f'border-radius:0.4rem;'
                    f'padding:0.6rem 0.8rem;'
                    f'margin:-0.25rem 0">'
                    f'<p style="margin:0 0 0.3rem 0;font-size:0.75rem;'
                    f'font-weight:700;color:{cfg["st_color"]};'
                    f'text-transform:uppercase;letter-spacing:0.05em">'
                    f'{cfg["emoji"]} {cfg["label"]}</p>'
                    f'{content}</div>',
                    unsafe_allow_html=True,
                )
        else:
            with st.chat_message("assistant"):
                st.markdown(content)


def extract_interrupt_prompt(value) -> str:
    """
    Extract a readable prompt from an interrupt payload.

    Supports:
    - plain strings
    - dicts like {"prompt": "..."}
    - LangGraph-like dicts such as {"value": "..."} or {"value": {"prompt": "..."}}
    """
    if isinstance(value, dict):
        raw_value = value.get("value", value)

        if isinstance(raw_value, dict):
            if "prompt" in raw_value:
                return str(raw_value["prompt"])
            return str(raw_value)

        return str(raw_value)

    return str(value)


def is_yes_no_interrupt(prompt: str | None) -> bool:
    """Return True if the interrupt prompt is marked as a Yes/No confirmation."""
    if not prompt:
        return False
    return YES_NO_MARKER in prompt


def clean_interrupt_prompt(prompt: str | None) -> str:
    """Remove the Yes/No marker before showing the prompt to the user."""
    if not prompt:
        return ""
    return prompt.replace(YES_NO_MARKER, "").strip()


def send_user_message(message: str, show_in_chat: bool = True):
    """Send a user message to the backend and process the response."""
    if show_in_chat:
        st.session_state.chat.append({"role": "user", "content": message})

    try:
        with st.spinner("Agents are thinking..."):
            resp = post_json(
                f"{st.session_state.server_url}/message/{st.session_state.thread_id}",
                {"message": message},
            )
        process_events(resp)
        time.sleep(0.05)
        st.rerun()
    except requests.ConnectionError:
        st.error(
            "Lost connection to the backend server. "
            "Check that it is still running at: "
            f"`{st.session_state.server_url}`"
        )
    except requests.RequestException as e:
        st.error(f"Failed to send message: {e}")


def process_events(resp: dict):
    """
    Process events from the backend response.

    resp:
      {
        "thread_id": "...",
        "events": [[node, key, value], ...],
        "needs_user_input": bool,
        "interrupt_prompt": str|None
      }
    """
    events = resp.get("events", [])
    found_interrupt = False

    for e in events:
        if not (isinstance(e, (list, tuple)) and len(e) == 3):
            continue

        node, key, value = e

        # Interrupt handling
        if key == "__interrupt__" or node == "__interrupt__":
            found_interrupt = True
            prompt = extract_interrupt_prompt(value)

            st.session_state.needs_user_input = True
            st.session_state.interrupt_prompt = prompt
            st.session_state.show_yes_no = is_yes_no_interrupt(prompt)
            continue

        # Avoid duplicate displays
        last_val = st.session_state.last_values.get(key)
        if last_val == value:
            continue
        st.session_state.last_values[key] = value

        # Track idea_board state in session
        if key == "idea_board" and value is not None:
            st.session_state.idea_board = str(value)

        if key == "final_file_name" and value:
            st.session_state.final_file_name = str(value)

        # Show download button when final_message is received
        if key == "final_message":
            st.session_state.show_download = True

        # Push displayable agent messages into chat log
        if should_display(str(key), value) or key == "final_message":
            role = role_from_key(str(node), str(key))

            # Force final message to come from facilitator
            if key == "final_message":
                role = "facilitator_structuring"

            st.session_state.chat.append(
                {"role": role, "content": str(value)}
            )

    # Fallback to top-level response fields if no interrupt event was found
    if not found_interrupt:
        st.session_state.needs_user_input = bool(resp.get("needs_user_input", False))

        fallback_prompt = resp.get("interrupt_prompt")
        if fallback_prompt:
            st.session_state.interrupt_prompt = str(fallback_prompt)
            st.session_state.show_yes_no = is_yes_no_interrupt(
                st.session_state.interrupt_prompt
            )
        else:
            st.session_state.interrupt_prompt = None
            st.session_state.show_yes_no = False

    # If the backend says no user input is needed anymore, clear interrupt UI state
    if not st.session_state.needs_user_input:
        st.session_state.interrupt_prompt = None
        st.session_state.show_yes_no = False


def render_homepage_guidelines():
    """Show homepage guidance explaining how to use the system."""
    with st.expander("Guidelines On Using The System", expanded=False):
        st.markdown(
            """
            This AI system works best when you treat the agents like a brainstorming team.
            The more specific and informative you are, the better the final essay plan will be.

            ### Guidelines for using the system
            
            - This system will help you brainstorm ideas to prepare you for your actual writing task.
            - The agents will not generate content for you. You create your own ideas, they are there to guide you.
            - Give detailed answers instead of short one-word replies.
            - Explain your ideas, even if they are rough or unfinished.
            - Share your opinion, examples, concerns, or possible arguments.
            - Tell the agents if you disagree with something or want a different direction.
            - When the agents ask questions, try to explain your reasoning rather than only saying yes or no.
            - Use the brainstorm to develop your own ideas, not just to wait for a finished answer.
            

            ### Sub-phases of the system

            This system consists of 3 main sub phases of brainstorming.

            **1. Idea exploration**  
            The agents ask questions to help you explore and expand on possible points, examples, and different angles.
            You will be talking with the Idea Coach and the Subject Specialist at this stage.

            **2. Critical review**  
            The critic challenges weak points, unclear reasoning, missing evidence, or one-sided arguments.
            You will be talking with the Critic at this stage.

            **3. Argument organisation**  
            The system helps arrange your ideas into a clearer essay structure with a logical flow.
            You will be talking with the Structure Coach and the Argument Flow Coach at this stage.

            ### What the main agents do

            **Idea Coach**  
            Helps you generate and clarify your initial ideas. This agent is useful when you are unsure what to say or need help developing rough thoughts.

            **Subject Specialist**  
            Helps to expand the ideas with content connected to the subject area. This agent can suggest more relevant concepts, examples, and academic angles.

            **Critic**  
            Challenges your ideas and points out weaknesses. This helps make your essay plan stronger by identifying vague claims, missing explanations, or unsupported arguments.

            **Structure coach**  
            Turns scattered ideas into clearer sections. This agent helps organise your ideas into a structured essay plan.

            **Argument Flow Coach**  
            Focuses on the order and logic of your essay. This agent helps make sure your introduction, main points, counterpoints, and conclusion connect properly.

            ### Remember

            - Only certain agents will talk to you at a phase. This makes sure you can effectively interact with all agents and systematically develop your essay plan.
            - You can skip to the next phase by saying : "Lets move on to the next phase" or something similar.
            - Once you exit a phase you can't re-enter it.
            - The system is not just asking questions for no reason. Each question is meant to help you build a stronger, clearer, and more organised essay plan.
            - SO ARE YOU READY TO GET BRAINSTORMING? GOOD LUCK !!


            """
        )


# ----------------------------
# File export helpers
# ----------------------------

def make_txt_file(content: str) -> str:
    """Return plain text content for TXT download."""
    return content


def make_docx_file(content: str) -> bytes:
    """Create a Word document in memory and return it as bytes."""
    buffer = BytesIO()

    doc = Document()
    doc.add_heading("Final Essay Plan", level=1)

    for line in content.splitlines():
        clean_line = line.strip()

        if not clean_line:
            doc.add_paragraph("")
        elif clean_line.startswith("# "):
            doc.add_heading(clean_line.replace("# ", "", 1), level=1)
        elif clean_line.startswith("## "):
            doc.add_heading(clean_line.replace("## ", "", 1), level=2)
        elif clean_line.startswith("### "):
            doc.add_heading(clean_line.replace("### ", "", 1), level=3)
        elif clean_line.startswith("- "):
            doc.add_paragraph(clean_line.replace("- ", "", 1), style="List Bullet")
        elif clean_line.startswith("* "):
            doc.add_paragraph(clean_line.replace("* ", "", 1), style="List Bullet")
        else:
            doc.add_paragraph(clean_line)

    doc.save(buffer)
    buffer.seek(0)

    return buffer.getvalue()


def make_pdf_safe_text(text: str) -> str:
    """
    Convert text to a PDF-safe form for FPDF core fonts.

    FPDF's built-in fonts do not support every Unicode character,
    so unsupported characters are replaced instead of crashing the app.
    """
    return text.encode("latin-1", "replace").decode("latin-1")


def break_long_words(text: str, max_word_length: int = 55) -> str:
    """
    Add spaces inside very long unbroken words.

    This prevents FPDF from failing on long URLs, markdown separators,
    or other text that cannot naturally wrap.
    """
    def split_word(match):
        word = match.group(0)
        return " ".join(
            word[i:i + max_word_length]
            for i in range(0, len(word), max_word_length)
        )

    return re.sub(r"\S{" + str(max_word_length + 1) + r",}", split_word, text)


def pdf_multi_cell_safe(pdf: FPDF, width: float, height: float, text: str):
    """
    Render text safely using explicit width and character wrapping where available.
    """
    pdf.set_x(pdf.l_margin)

    try:
        pdf.multi_cell(
            width,
            height,
            text,
            new_x="LMARGIN",
            new_y="NEXT",
            wrapmode="CHAR",
        )
    except TypeError:
        # Older fpdf2 versions may not support wrapmode.
        pdf.set_x(pdf.l_margin)
        pdf.multi_cell(
            width,
            height,
            break_long_words(text),
            new_x="LMARGIN",
            new_y="NEXT",
        )
    except FPDFException:
        # Final fallback: shrink font slightly and forcibly split long words.
        pdf.set_font("Helvetica", "", 9)
        pdf.set_x(pdf.l_margin)
        pdf.multi_cell(
            width,
            height,
            break_long_words(text, max_word_length=35),
            new_x="LMARGIN",
            new_y="NEXT",
        )


def make_pdf_file(content: str) -> bytes:
    """Create a simple PDF in memory and return it as bytes."""
    pdf = FPDF()
    pdf.set_margins(left=15, top=15, right=15)
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()

    usable_width = pdf.epw

    pdf.set_font("Helvetica", "B", 16)
    pdf_multi_cell_safe(pdf, usable_width, 10, "Final Essay Plan")

    pdf.ln(4)

    for line in content.splitlines():
        clean_line = make_pdf_safe_text(line.strip())

        if not clean_line:
            pdf.ln(4)
            continue

        if clean_line.startswith("# "):
            text = clean_line.replace("# ", "", 1)
            pdf.set_font("Helvetica", "B", 15)
        elif clean_line.startswith("## "):
            text = clean_line.replace("## ", "", 1)
            pdf.set_font("Helvetica", "B", 13)
        elif clean_line.startswith("### "):
            text = clean_line.replace("### ", "", 1)
            pdf.set_font("Helvetica", "B", 12)
        elif clean_line.startswith("- "):
            text = "- " + clean_line.replace("- ", "", 1)
            pdf.set_font("Helvetica", "", 11)
        elif clean_line.startswith("* "):
            text = "- " + clean_line.replace("* ", "", 1)
            pdf.set_font("Helvetica", "", 11)
        else:
            text = clean_line
            pdf.set_font("Helvetica", "", 11)

        text = break_long_words(text)
        pdf_multi_cell_safe(pdf, usable_width, 7, text)
        pdf.ln(1)

    raw_pdf = pdf.output()

    if isinstance(raw_pdf, str):
        return raw_pdf.encode("latin-1")

    return bytes(raw_pdf)


def build_download_file(content: str, file_type: str, base_name: str):
    """Return data, filename, and MIME type for selected download format."""
    safe_base_name = Path(base_name).stem or "essay_plan"

    if file_type == "TXT":
        return (
            make_txt_file(content),
            f"{safe_base_name}.txt",
            "text/plain",
        )

    if file_type == "Word document":
        return (
            make_docx_file(content),
            f"{safe_base_name}.docx",
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )

    if file_type == "PDF":
        return (
            make_pdf_file(content),
            f"{safe_base_name}.pdf",
            "application/pdf",
        )

    return (
        make_txt_file(content),
        f"{safe_base_name}.txt",
        "text/plain",
    )


# ----------------------------
# Session state init
# ----------------------------

if "server_url" not in st.session_state:
    st.session_state.server_url = os.environ.get("BACKEND_URL", "http://127.0.0.1:8000")
if "thread_id" not in st.session_state:
    st.session_state.thread_id = None
if "chat" not in st.session_state:
    st.session_state.chat = []
if "last_values" not in st.session_state:
    st.session_state.last_values = {}
if "needs_user_input" not in st.session_state:
    st.session_state.needs_user_input = False
if "interrupt_prompt" not in st.session_state:
    st.session_state.interrupt_prompt = None
if "show_yes_no" not in st.session_state:
    st.session_state.show_yes_no = False
if "idea_board" not in st.session_state:
    st.session_state.idea_board = None
if "show_download" not in st.session_state:
    st.session_state.show_download = False
if "final_file_name" not in st.session_state:
    st.session_state.final_file_name = "essay_plan.txt"


# ----------------------------
# UI
# ----------------------------

st.title("EssayPlanner")

with st.sidebar:
    st.subheader("Server")
    st.session_state.server_url = st.text_input(
        "FastAPI base URL", st.session_state.server_url
    )
    st.caption("Backend should expose: POST /start and POST /message/{thread_id}")

    st.markdown("---")

    # Agent legend
    with st.expander("Agent Roles", expanded=True):
        st.markdown("🧑‍🎓 **You** — your messages")
        for cfg in AGENT_CONFIG.values():
            swatch = (
                f'<span style="display:inline-block;width:12px;height:12px;'
                f'border-radius:3px;background:{cfg["bg_color"]};'
                f'border:2px solid {cfg["st_color"]};vertical-align:middle;'
                f'margin-right:6px;"></span>'
            )
            st.markdown(
                f'{swatch}{cfg["emoji"]} **{cfg["label"]}**',
                unsafe_allow_html=True,
            )

    st.markdown("---")

    # Idea Board
    with st.expander("Idea Board", expanded=True):
        if st.session_state.idea_board:
            st.markdown(st.session_state.idea_board)
        else:
            st.caption("No ideas yet. Start a session to begin.")

    st.markdown("---")

    if st.button("Reset session"):
        st.session_state.thread_id = None
        st.session_state.chat = []
        st.session_state.last_values = {}
        st.session_state.needs_user_input = False
        st.session_state.interrupt_prompt = None
        st.session_state.show_yes_no = False
        st.session_state.idea_board = None
        st.session_state.show_download = False
        st.session_state.final_file_name = "essay_plan.txt"
        st.rerun()


# Render chat history
for msg in st.session_state.chat:
    render_chat_message(msg["role"], msg["content"])


# ----------------------------
# Start flow (no session yet)
# ----------------------------

if st.session_state.thread_id is None:
    st.subheader("Start a new session")

    render_homepage_guidelines()

    subject = st.text_input("Subject", value="Education")
    essay_topic = st.text_input("Essay topic", value="AI in education and creativity")

    if st.button("Start"):
        try:
            with st.spinner("Agents are thinking..."):
                resp = post_json(
                    f"{st.session_state.server_url}/start",
                    {"subject": subject, "essay_topic": essay_topic},
                )
            st.session_state.thread_id = resp["thread_id"]
            process_events(resp)
            st.rerun()
        except requests.ConnectionError:
            st.error(
                "Could not connect to the backend server. "
                "Make sure it is running at: "
                f"`{st.session_state.server_url}`"
            )
            st.info(
                "Start the backend with:\n\n"
                "```\n"
                "uv run uvicorn bin.MultiAgentSystem.app:app "
                "--reload --host 127.0.0.1 --port 8000\n"
                "```"
            )
        except requests.RequestException as e:
            st.error(f"Failed to start session: {e}")

    st.stop()


# ----------------------------
# Active session flow
# ----------------------------

# Show Yes/No buttons only for prompts marked with [YES_NO]
if st.session_state.show_yes_no and st.session_state.interrupt_prompt:
    st.info(clean_interrupt_prompt(st.session_state.interrupt_prompt))

    col1, col2 = st.columns(2)

    if col1.button("Yes", use_container_width=True, key="yes_button"):
        st.session_state.show_yes_no = False
        st.session_state.interrupt_prompt = None
        send_user_message("yes", show_in_chat=False)

    if col2.button("No", use_container_width=True, key="no_button"):
        st.session_state.show_yes_no = False
        st.session_state.interrupt_prompt = None
        send_user_message("no", show_in_chat=False)

    # IMPORTANT FIX:
    # Keep the pinned chat input rendered even when Yes/No buttons appear.
    # Without this, st.stop() prevents the normal bottom chat input from rendering,
    # which changes the page layout and can make the view jump upward.
    st.chat_input(
        "Choose Yes or No above",
        disabled=True,
        key="yes_no_disabled_chat_input",
    )

    st.stop()


# For other interrupts, show the prompt and keep normal text input available
if (
    st.session_state.needs_user_input
    and st.session_state.interrupt_prompt
    and not st.session_state.show_yes_no
):
    st.info(clean_interrupt_prompt(st.session_state.interrupt_prompt))


user_text = st.chat_input("Type your reply here", key="normal_chat_input")

if user_text:
    send_user_message(user_text)


# ----------------------------
# Final download section
# ----------------------------

if st.session_state.show_download and st.session_state.idea_board:
    st.markdown("---")
    st.success("Your essay plan is ready!")

    st.markdown("### 📄 Final Essay Plan")

    file_type = st.radio(
        "Choose download format",
        ["TXT", "Word document", "PDF"],
        horizontal=True,
        key="download_file_type",
    )

    file_data, file_name, mime_type = build_download_file(
        content=st.session_state.idea_board,
        file_type=file_type,
        base_name=st.session_state.final_file_name,
    )

    st.download_button(
        label=f"⬇️ Download as {file_type}",
        data=file_data,
        file_name=file_name,
        mime=mime_type,
    )
